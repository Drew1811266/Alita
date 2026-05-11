from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document


SUPPORTED_INPUTS = {".txt", ".md", ".docx"}


@dataclass(frozen=True)
class DocumentReadResult:
    text: str
    sources: list[str]


def read_documents(paths: list[str]) -> DocumentReadResult:
    chunks: list[str] = []
    sources: list[str] = []

    for raw_path in paths:
        path = Path(raw_path)
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_INPUTS:
            raise ValueError(f"unsupported_format:{suffix}")

        if suffix in {".txt", ".md"}:
            text = _read_text_file(path)
        else:
            text = _read_docx_file(path)

        chunks.append(text)
        sources.append(str(path))

    return DocumentReadResult(text="\n\n".join(chunks), sources=sources)


def write_markdown(content: str, output_path: str) -> str:
    path = Path(output_path)
    if path.suffix.lower() != ".md":
        raise ValueError("write_failed:markdown_output_must_end_with_md")

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return str(path)


def write_docx(content: str, output_path: str) -> str:
    path = Path(output_path)
    if path.suffix.lower() != ".docx":
        raise ValueError("write_failed:docx_output_must_end_with_docx")

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for line in content.splitlines():
        if line.strip():
            document.add_paragraph(line)
    document.save(path)
    return str(path)


def _read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except OSError as error:
        raise ValueError(f"read_failed:{path}") from error


def _read_docx_file(path: Path) -> str:
    try:
        document = Document(path)
    except Exception as error:
        raise ValueError(f"read_failed:{path}") from error

    return "\n".join(
        paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()
    )
