from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from agent_service.execution import run_graph_events
from tests.test_execution import (
    FakeModelClient,
    FakeToolExecutor,
    TypstFlowToolExecutor,
    build_document_flow_request,
    build_document_flow_request_with_typst,
)
from tools.document_tool import read_documents, write_docx


def test_markdown_document_fixture_exports_report_artifact(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# 标题\n\n这是一段中文正文。", encoding="utf-8")
    request = build_document_flow_request(tmp_path, source)

    events = list(
        run_graph_events(
            request,
            model_client=FakeModelClient(),
            tool_executor=FakeToolExecutor(),
        )
    )

    file_export_events = [
        event
        for event in events
        if event.type == "artifact.created"
        and event.payload["sourceNodeId"] == "file-export"
    ]
    assert len(file_export_events) == 1
    artifact_path = Path(file_export_events[0].payload["path"])
    assert artifact_path.is_file()
    assert artifact_path.suffix == ".md"

    content = artifact_path.read_text(encoding="utf-8")
    assert "outline result" in content
    assert "report result" in content


def test_docx_fixture_can_be_read_and_exported(tmp_path: Path) -> None:
    source = tmp_path / "input.docx"
    document = Document()
    document.add_paragraph("DOCX fixture heading")
    document.add_paragraph("包含中文的段落")
    document.save(source)

    result = read_documents([str(source)])

    assert "DOCX fixture heading" in result.text
    assert "包含中文的段落" in result.text
    assert result.sources == [str(source)]

    exported_path = Path(
        write_docx("DOCX fixture heading\n\n包含中文的段落", str(tmp_path / "export.docx"))
    )
    exported = Document(exported_path)
    assert [paragraph.text for paragraph in exported.paragraphs] == [
        "DOCX fixture heading",
        "包含中文的段落",
    ]


def test_corrupt_docx_fixture_reports_clear_read_failure(tmp_path: Path) -> None:
    source = tmp_path / "corrupt.docx"
    source.write_bytes(b"not a valid docx archive")

    with pytest.raises(ValueError, match="read_failed:"):
        read_documents([str(source)])


def test_typst_flow_exports_pdf_artifact_path(tmp_path: Path) -> None:
    source = tmp_path / "source.md"
    source.write_text("# 标题\n\n这是一段中文正文。", encoding="utf-8")
    request = build_document_flow_request_with_typst(tmp_path, source)
    tool_executor = TypstFlowToolExecutor()

    events = list(
        run_graph_events(
            request,
            model_client=FakeModelClient(),
            tool_executor=tool_executor,
        )
    )

    file_export_events = [
        event
        for event in events
        if event.type == "artifact.created"
        and event.payload["sourceNodeId"] == "file-export"
    ]
    pdf_artifact = next(
        Path(event.payload["path"])
        for event in file_export_events
        if Path(event.payload["path"]).suffix == ".pdf"
    )
    assert pdf_artifact.is_file()
    assert pdf_artifact.read_bytes().startswith(b"%PDF")
