from pathlib import Path

import pytest
from docx import Document

from tools.document_tool import read_documents, write_docx, write_markdown


def test_reads_txt_and_markdown(tmp_path: Path) -> None:
    txt = tmp_path / "a.txt"
    md = tmp_path / "b.md"
    txt.write_text("hello", encoding="utf-8")
    md.write_text("# title", encoding="utf-8")

    result = read_documents([str(txt), str(md)])

    assert "hello" in result.text
    assert "# title" in result.text
    assert result.sources == [str(txt), str(md)]


def test_reads_docx(tmp_path: Path) -> None:
    path = tmp_path / "input.docx"
    doc = Document()
    doc.add_paragraph("docx content")
    doc.save(path)

    result = read_documents([str(path)])

    assert "docx content" in result.text
    assert result.sources == [str(path)]


def test_writes_markdown_and_docx(tmp_path: Path) -> None:
    md_path = write_markdown("report", str(tmp_path / "report.md"))
    docx_path = write_docx("report\n\nsecond paragraph", str(tmp_path / "report.docx"))

    assert Path(md_path).read_text(encoding="utf-8") == "report"
    assert Path(docx_path).exists()

    exported = Document(docx_path)
    assert [paragraph.text for paragraph in exported.paragraphs] == [
        "report",
        "second paragraph",
    ]


def test_rejects_unsupported_input_format(tmp_path: Path) -> None:
    pdf = tmp_path / "input.pdf"
    pdf.write_bytes(b"%PDF")

    with pytest.raises(ValueError, match="unsupported_format:.pdf"):
        read_documents([str(pdf)])


def test_rejects_wrong_output_suffix(tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="markdown_output_must_end_with_md"):
        write_markdown("report", str(tmp_path / "report.txt"))

    with pytest.raises(ValueError, match="docx_output_must_end_with_docx"):
        write_docx("report", str(tmp_path / "report.md"))
